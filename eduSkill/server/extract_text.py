import sys
import os
import subprocess


def extract(path):
    if not os.path.exists(path):
        raise FileNotFoundError(f"Không tìm thấy tệp: {path}")

    ext = os.path.splitext(path)[1].lower()
    if ext in ('.txt', '.md', '.tex', '.csv'):
        return open(path, encoding='utf-8', errors='replace').read()
    if ext == '.docx':
        from docx import Document
        d = Document(path)
        parts = [p.text for p in d.paragraphs]
        for tb in d.tables:
            for r in tb.rows:
                parts.append(' | '.join(c.text for c in r.cells))
        return '\n'.join(x for x in parts if x is not None)
    if ext == '.pdf':
        out = path + '.txt'
        try:
            subprocess.run(['pdftotext', '-layout', path, out], check=True, capture_output=True)
            t = open(out, encoding='utf-8', errors='replace').read()
            try:
                os.remove(out)
            except OSError:
                pass
            return t
        except Exception as e1:
            # Dự phòng cho máy không có pdftotext hoặc PDF khiến poppler lỗi.
            try:
                import fitz  # pymupdf
                doc = fitz.open(path)
                return '\n'.join(p.get_text() for p in doc)
            except Exception as e2:
                raise RuntimeError(f"Không đọc được PDF. pdftotext: {e1}; PyMuPDF: {e2}")
    # Mặc định: đọc như văn bản thuần
    return open(path, encoding='utf-8', errors='replace').read()


if __name__ == "__main__":
    sys.stdout.write(extract(sys.argv[1]))
