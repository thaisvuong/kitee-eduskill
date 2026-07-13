# -*- coding: utf-8 -*-
"""
Kitee_Word_Template.py
Tạo file Word template đơn giản, chỉnh chu cho phiếu học/giáo trình KiTee.

Yêu cầu chính:
- Vào thẳng phần Lý thuyết cần nhớ.
- Không gom nhiều câu hỏi vào một box.
- Mỗi câu/bài tập là một box riêng.
- Có chừa khoảng trống làm bài trong từng box.
- Có phần Lưu ý định dạng ở cuối file.

Cách chạy:
    pip install python-docx
    python Kitee_Word_Template.py

File xuất ra:
    Kitee_Word_Template.docx
"""

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =========================
# 1) Cấu hình thương hiệu
# =========================
OUTPUT_FILE = "/mnt/data/Kitee_Word_Template.docx"

BRAND_NAVY = "1B3C6E"
BRAND_ORANGE = "E8741C"
TEXT_DARK = "1F2937"
TEXT_MUTED = "667085"
BORDER = "D8E2F0"
LIGHT_BLUE = "F4F8FD"
LIGHT_ORANGE = "FFF4EC"
LIGHT_GRAY = "F8FAFC"
WHITE = "FFFFFF"

FONT_BODY = "Times New Roman"
FONT_HEAD = "Arial"


# =========================
# 2) Helper định dạng cơ bản
# =========================
def set_run_font(run, font=FONT_BODY, size=12.5, color=TEXT_DARK, bold=False, italic=False):
    """Định dạng font cho run, hỗ trợ tiếng Việt ổn định."""
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:eastAsia"), font)
    r_fonts.set(qn("w:cs"), font)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_paragraph_border(paragraph, color=BORDER, size="8", space="3", sides=("bottom",)):
    """Thêm đường kẻ cho paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for side in sides:
        tag = f"w:{side}"
        element = p_bdr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            p_bdr.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), space)
        element.set(qn("w:color"), color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="10"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=140, start=180, bottom=140, end=180):
    """Cell margin theo twips. 1 cm khoảng 567 twips."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width_pct(table, pct=100):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(pct * 50))


def prevent_table_row_split(table):
    """Giữ mỗi box không bị tách đôi qua 2 trang khi render/in."""
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


# =========================
# 3) Thiết lập tài liệu
# =========================
def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(12.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT_DARK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)

    for style_name in ["Heading 1", "Heading 2"]:
        st = doc.styles[style_name]
        st.font.name = FONT_HEAD
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(BRAND_NAVY)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEAD)

    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 2"].font.size = Pt(14)

    if "KiteeCaption" not in [s.name for s in doc.styles]:
        st = doc.styles.add_style("KiteeCaption", WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = FONT_HEAD
        st.font.size = Pt(9)
        st.font.color.rgb = RGBColor.from_string(TEXT_MUTED)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEAD)

    return doc


def add_header_footer(doc):
    section = doc.sections[0]

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.text = ""
    r = header.add_run("KiTee Education | Phiếu học")
    set_run_font(r, font=FONT_HEAD, size=9, color=TEXT_MUTED, bold=True)
    set_paragraph_border(header, color=BORDER, size="4", sides=("bottom",))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = ""
    r = footer.add_run("Kitee_Word_Template | Dùng để tạo giáo trình, phiếu học, tài liệu ôn tập")
    set_run_font(r, font=FONT_HEAD, size=8.5, color=TEXT_MUTED)


# =========================
# 4) Component nội dung
# =========================
def add_document_title(doc, title, subtitle=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KITEE WORD TEMPLATE")
    set_run_font(r, font=FONT_HEAD, size=10, color=BRAND_ORANGE, bold=True)
    set_paragraph_spacing(p, before=4, after=4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, font=FONT_HEAD, size=22, color=BRAND_NAVY, bold=True)
    set_paragraph_spacing(p, before=0, after=2)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        set_run_font(r, font=FONT_HEAD, size=10.5, color=TEXT_MUTED)
        set_paragraph_spacing(p, before=0, after=8)

    line = doc.add_paragraph()
    set_paragraph_border(line, color=BRAND_ORANGE, size="12", sides=("bottom",))
    set_paragraph_spacing(line, before=0, after=10)


def add_section_title(doc, number, title):
    p = doc.add_paragraph()
    r = p.add_run(f"{number}. ")
    set_run_font(r, font=FONT_HEAD, size=15.5, color=BRAND_ORANGE, bold=True)
    r = p.add_run(title)
    set_run_font(r, font=FONT_HEAD, size=15.5, color=BRAND_NAVY, bold=True)
    set_paragraph_spacing(p, before=10, after=4)
    set_paragraph_border(p, color=BORDER, size="6", sides=("bottom",))


def add_text(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, font=FONT_BODY, size=12.5, color=TEXT_DARK, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, font=FONT_BODY, size=12.5, color=TEXT_DARK)
    else:
        r = p.add_run(text)
        set_run_font(r, font=FONT_BODY, size=12.5, color=TEXT_DARK)
    set_paragraph_spacing(p, after=5)
    return p


def add_bullet(doc, text, color=BRAND_ORANGE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    r = p.add_run("• ")
    set_run_font(r, font=FONT_HEAD, size=12.5, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, font=FONT_BODY, size=12.5, color=TEXT_DARK)
    set_paragraph_spacing(p, after=3)
    return p


def add_box(doc, label, body_lines, fill=LIGHT_BLUE, border=BRAND_NAVY, label_color=BRAND_NAVY):
    """Tạo một box riêng. body_lines là list[str]."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    set_table_width_pct(table, 100)
    prevent_table_row_split(table)

    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=border, size="9")
    set_cell_margins(cell, top=130, start=190, bottom=130, end=190)

    # Xoá paragraph rỗng mặc định
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    r = p.add_run(label.upper())
    set_run_font(r, font=FONT_HEAD, size=10.5, color=label_color, bold=True)
    set_paragraph_spacing(p, after=4)

    for line in body_lines:
        p = cell.add_paragraph()
        r = p.add_run(line)
        set_run_font(r, font=FONT_BODY, size=12.4, color=TEXT_DARK)
        set_paragraph_spacing(p, after=3)

    # Khoảng cách sau box
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4)
    return table


def add_theory_box(doc, label, lines):
    return add_box(doc, label, lines, fill=LIGHT_BLUE, border=BRAND_NAVY, label_color=BRAND_NAVY)


def add_example_box(doc, label, lines):
    return add_box(doc, label, lines, fill=LIGHT_ORANGE, border=BRAND_ORANGE, label_color=BRAND_ORANGE)


def add_answer_lines(cell, count=5):
    """Thêm dòng chừa chỗ làm bài trong cell của box câu hỏi."""
    p = cell.add_paragraph()
    r = p.add_run("Lời làm:")
    set_run_font(r, font=FONT_HEAD, size=10.2, color=BRAND_NAVY, bold=True)
    set_paragraph_spacing(p, before=4, after=2)

    for _ in range(count):
        p = cell.add_paragraph()
        r = p.add_run("................................................................................................................")
        set_run_font(r, font=FONT_BODY, size=11.5, color=TEXT_MUTED)
        set_paragraph_spacing(p, before=0, after=1, line=1.0)


def add_question_box(doc, number, question, work_lines=5, hint=None):
    """Mỗi câu/bài tập là một box riêng, có sẵn chỗ làm bài."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width_pct(table, 100)
    prevent_table_row_split(table)

    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_border(cell, color=BORDER, size="9")
    set_cell_margins(cell, top=150, start=190, bottom=150, end=190)

    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    r = p.add_run(f"BÀI {number}")
    set_run_font(r, font=FONT_HEAD, size=10.5, color=BRAND_ORANGE, bold=True)
    set_paragraph_spacing(p, after=3)

    p = cell.add_paragraph()
    r = p.add_run(question)
    set_run_font(r, font=FONT_BODY, size=12.5, color=TEXT_DARK)
    set_paragraph_spacing(p, after=3)

    if hint:
        p = cell.add_paragraph()
        r = p.add_run(f"Gợi ý: {hint}")
        set_run_font(r, font=FONT_BODY, size=11.5, color=TEXT_MUTED, italic=True)
        set_paragraph_spacing(p, after=2)

    add_answer_lines(cell, count=work_lines)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4)
    return table


def add_format_note_box(doc):
    """Phần lưu ý định dạng đặt ở cuối file."""
    add_section_title(doc, "6", "Lưu ý định dạng")
    add_box(
        doc,
        "Lưu ý khi dùng template",
        [
            "1. Mỗi câu hỏi/bài tập đặt trong một box riêng, không gom nhiều câu vào cùng một box.",
            "2. Sau mỗi câu hỏi luôn chừa phần 'Lời làm' với 4–6 dòng trống để học sinh làm bài.",
            "3. Tiêu đề chính dùng Arial Bold, màu navy #1B3C6E; điểm nhấn dùng cam #E8741C.",
            "4. Nội dung bài học dùng Times New Roman 12.5pt, giãn dòng 1.15 để dễ đọc và dễ in.",
            "5. Box lý thuyết dùng nền xanh nhạt; box ví dụ dùng nền cam nhạt; box bài tập dùng nền xám rất nhạt.",
            "6. Khi thêm bài mới trong Python, dùng hàm add_question_box(doc, number, question, work_lines, hint).",
        ],
        fill=WHITE,
        border=BRAND_ORANGE,
        label_color=BRAND_ORANGE,
    )


# =========================
# 5) Nội dung mẫu
# =========================
def build_document():
    doc = setup_document()
    add_header_footer(doc)

    add_document_title(
        doc,
        "BÀI 01. LÀM QUEN VỚI PHƯƠNG TRÌNH",
        "Lý thuyết cần nhớ • Ví dụ mẫu • Bài tập riêng từng box • Chừa chỗ làm bài",
    )

    # Vào thẳng lý thuyết cần nhớ
    add_section_title(doc, "1", "Lý thuyết cần nhớ")
    add_text(
        doc,
        "Phương trình là một đẳng thức có chứa ẩn số. Khi giải phương trình, ta cần tìm giá trị của ẩn để đẳng thức trở thành đúng.",
    )
    add_theory_box(
        doc,
        "Ghi nhớ 1",
        [
            "Muốn tìm số hạng chưa biết trong phép cộng, ta lấy tổng trừ đi số hạng đã biết.",
            "Dạng mẫu: x + a = b  →  x = b - a.",
        ],
    )
    add_theory_box(
        doc,
        "Ghi nhớ 2",
        [
            "Muốn tìm số bị trừ, ta lấy hiệu cộng với số trừ.",
            "Dạng mẫu: x - a = b  →  x = b + a.",
        ],
    )
    add_theory_box(
        doc,
        "Ghi nhớ 3",
        [
            "Muốn tìm thừa số chưa biết, ta lấy tích chia cho thừa số đã biết.",
            "Dạng mẫu: a × x = b  →  x = b : a.",
        ],
    )

    add_section_title(doc, "2", "Cách trình bày mẫu")
    add_bullet(doc, "Bước 1: Xác định phép tính đang xuất hiện trong phương trình.")
    add_bullet(doc, "Bước 2: Dùng phép tính ngược để tìm giá trị của ẩn.")
    add_bullet(doc, "Bước 3: Kết luận nghiệm và thay lại vào đề để kiểm tra.")

    doc.add_page_break()
    add_section_title(doc, "3", "Ví dụ minh họa")
    add_example_box(
        doc,
        "Ví dụ 1",
        [
            "Giải phương trình: x + 7 = 15",
            "Ta có: x = 15 - 7 = 8.",
            "Vậy nghiệm của phương trình là x = 8.",
        ],
    )
    add_example_box(
        doc,
        "Ví dụ 2",
        [
            "Giải phương trình: x - 5 = 13",
            "Ta có: x = 13 + 5 = 18.",
            "Vậy nghiệm của phương trình là x = 18.",
        ],
    )

    doc.add_page_break()
    add_section_title(doc, "4", "Bài tập tự luyện")

    # Mỗi câu là một box riêng, có chỗ làm bài
    add_question_box(doc, "1", "Giải phương trình: x + 9 = 20.", work_lines=5)
    add_question_box(doc, "2", "Giải phương trình: x - 6 = 14.", work_lines=5)
    add_question_box(doc, "3", "Giải phương trình: 3 × x = 24.", work_lines=5)

    doc.add_page_break()
    add_section_title(doc, "4", "Bài tập tự luyện (tiếp theo)")
    add_question_box(
        doc,
        "4",
        "Một số cộng với 12 thì bằng 35. Hỏi số đó là bao nhiêu? Hãy viết lời giải bằng một phương trình đơn giản.",
        work_lines=6,
    )
    add_question_box(
        doc,
        "5",
        "Tìm x biết: 2 × x + 5 = 21.",
        work_lines=6,
        hint="Xử lý phép cộng trước, sau đó xử lý phép nhân.",
    )

    doc.add_page_break()
    add_section_title(doc, "5", "Lời giải tham khảo")
    add_box(doc, "Bài 1", ["x + 9 = 20  →  x = 20 - 9 = 11."], fill=WHITE, border=BORDER, label_color=BRAND_NAVY)
    add_box(doc, "Bài 2", ["x - 6 = 14  →  x = 14 + 6 = 20."], fill=WHITE, border=BORDER, label_color=BRAND_NAVY)
    add_box(doc, "Bài 3", ["3 × x = 24  →  x = 24 : 3 = 8."], fill=WHITE, border=BORDER, label_color=BRAND_NAVY)
    add_box(
        doc,
        "Bài 4",
        ["Gọi số cần tìm là x. Theo đề bài: x + 12 = 35.", "Suy ra: x = 35 - 12 = 23. Vậy số cần tìm là 23."],
        fill=WHITE,
        border=BORDER,
        label_color=BRAND_NAVY,
    )
    add_box(
        doc,
        "Bài 5",
        ["2 × x + 5 = 21  →  2 × x = 21 - 5 = 16.", "Suy ra: x = 16 : 2 = 8. Vậy x = 8."],
        fill=WHITE,
        border=BORDER,
        label_color=BRAND_NAVY,
    )

    add_format_note_box(doc)

    doc.save(OUTPUT_FILE)
    print(f"Đã tạo file: {OUTPUT_FILE}")


if __name__ == "__main__":
    build_document()
