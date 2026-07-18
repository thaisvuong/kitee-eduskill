import sys
import json
import os
import re
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

# Chuyển LaTeX -> OMML (chế độ Công thức của Word). Nếu thiếu lib -> fallback text.
try:
  import latex2mathml.converter as _l2m
  from mathml2omml import convert as _m2o
  _MATH_OK = True
except Exception:
  _MATH_OK = False

_OMML_NS = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'

# Chiều rộng vùng nội dung (cm) — tính lại trong create_styled_word theo khổ giấy/lề.
CONTENT_WIDTH_CM = 17.0

# ────────────────────────────── Brand ──────────────────────────────
BRAND_NAVY = "1B3C6E"
BRAND_ORANGE = "E8741C"
TEXT_DARK = "1F2937"
TEXT_MUTED = "667085"
BORDER = "D8E2F0"
LIGHT_BLUE = "F4F8FD"
LIGHT_ORANGE = "FFF4EC"
LIGHT_GRAY = "F8FAFC"
WHITE = "FFFFFF"

FONT_BODY = "Times New Roman"
FONT_HEAD = "Arial"

BODY_SIZE = 12.5

# ─────────────────────────── Text helpers ──────────────────────────
def normalize_ws(text):
  """Chuyển escape literal còn sót (\\n, \\t, \\r do JSON/model để lại) thành
  khoảng trắng thật — NHƯNG chỉ ở phần văn xuôi, GIỮ NGUYÊN vùng công thức $...$
  (để không phá \\times, \\text, \\frac bên trong)."""
  if not text:
    return text
  out = []
  for seg in re.split(r'(\$[^$]+\$)', text):
    if seg.startswith('$') and seg.endswith('$'):
      out.append(seg)           # giữ nguyên math
    else:
      out.append(seg.replace('\\r\\n', '\n').replace('\\n', '\n')
             .replace('\\r', '\n').replace('\\t', ' '))
  return ''.join(out)


def strip_markdown(text):
  """Xoá toàn bộ ký tự Markdown (dùng cho nhãn / tiêu đề ngắn)."""
  if not text:
    return ""
  text = normalize_ws(text)
  text = re.sub(r'#+\s*', '', text)
  text = text.replace('**', '').replace('__', '').replace('`', '')
  text = re.sub(r'^\s*[\*\-\+]\s+', '', text, flags=re.MULTILINE)
  return text.strip()


def clean_caption(text):
  """Chú thích hình: bỏ markup toán ($...$, \\text, \\frac...) -> chữ thường gọn."""
  if not text:
    return ""
  text = re.sub(r'\\text\{([^}]*)\}', r'\1', text)
  text = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', text)
  text = re.sub(r'\\d?frac\{([^}]*)\}\{([^}]*)\}', r'\1/\2', text)
  text = text.replace('\\times', 'x').replace('\\,', ' ')
  text = re.sub(r'[$\\]', '', text)
  return strip_markdown(re.sub(r'\s+', ' ', text))


def add_math(paragraph, latex):
  """Chèn công thức dạng OMML (chế độ Công thức của Word) vào paragraph.

  Nếu chuyển đổi thất bại -> ghi lại dạng text để không mất nội dung.
  """
  latex = latex.strip()
  if _MATH_OK and latex:
    try:
      mml = _l2m.convert(latex)
      omml = _m2o(mml)
      # Chỉ lấy các khối <m:oMath>, gắn khai báo namespace để lxml phân tích được.
      for chunk in re.findall(r'<m:oMath>.*?</m:oMath>', omml, re.DOTALL):
        xml = chunk.replace('<m:oMath>', f'<m:oMath {_OMML_NS}>', 1)
        paragraph._p.append(parse_xml(xml))
      return
    except Exception:
      pass
  run = paragraph.add_run(latex)
  set_run_font(run, size=BODY_SIZE, italic=True)


def normalize_inline_math_units(text):
  """Sửa các dạng mũ đơn vị hay gặp trước khi chuyển sang OMML.

  LLM thường sinh `m$^2$` hoặc `cm$^3$`: phần cơ sở nằm ngoài công thức,
  còn `^2` nằm trong công thức. Word/OMML sẽ render thành ô vuông có mũ 2.
  Chuẩn hóa về `$m^2$`, `$cm^3$`; đồng thời xử lý Unicode `m²`, `cm³`.
  """
  if not text:
    return text
  # m$^2$, cm$^{3}$ -> $m^2$, $cm^3$
  text = re.sub(r'\b([A-Za-zÀ-ỹ]+)\s*\$\s*\^\s*\{?([23])\}?\s*\$', r'$\1^\2$', text)
  # m^2, cm^3 nằm ngoài công thức -> $m^2$, $cm^3$ (tránh bắt trong $...$ bằng cách chỉ xử lý text phổ thông)
  text = re.sub(r'\b(mm|cm|dm|m|km|ha)\s*\^\s*([23])\b', r'$\1^\2$', text, flags=re.IGNORECASE)
  # Unicode superscript hay gặp
  sup = {'²': '2', '³': '3'}
  text = re.sub(r'\b(mm|cm|dm|m|km|ha)([²³])\b', lambda m: f"${m.group(1)}^{sup[m.group(2)]}$", text, flags=re.IGNORECASE)
  return text


def add_inline_runs(paragraph, text, size=BODY_SIZE, color=TEXT_DARK,
          font=FONT_BODY, base_bold=False):
  """Thêm text vào paragraph: biên dịch $công thức$ -> OMML, **đậm**/*nghiêng* -> run.

  Đây là chốt chặn đảm bảo KHÔNG còn ký tự Markdown lọt vào file Word.
  """
  text = normalize_inline_math_units(text)
  text = re.sub(r'#+\s*', '', text)     # bỏ dấu heading còn sót
  text = text.replace('`', '')
  # 1) Tách công thức $...$ trước
  for seg in re.split(r'(\$[^$]+\$)', text):
    if not seg:
      continue
    if seg.startswith('$') and seg.endswith('$') and len(seg) > 2:
      add_math(paragraph, seg[1:-1])
      continue
    # 2) Trong phần chữ thường: tách **đậm** / *nghiêng*
    for part in re.split(r'(\*\*[^*]+\*\*|\*[^*\n]+\*)', seg):
      if not part:
        continue
      bold = base_bold
      italic = False
      content = part
      if part.startswith('**') and part.endswith('**') and len(part) > 4:
        bold = True
        content = part[2:-2]
      elif part.startswith('*') and part.endswith('*') and len(part) > 2:
        italic = True
        content = part[1:-1]
      content = content.replace('*', '')   # dọn dấu * lẻ
      if not content:
        continue
      run = paragraph.add_run(content)
      set_run_font(run, font=font, size=size, color=color, bold=bold, italic=italic)


def set_run_font(run, font=FONT_BODY, size=BODY_SIZE, color=TEXT_DARK, bold=False, italic=False):
  run.font.name = font
  run.font.size = Pt(size)
  run.font.bold = bold
  run.font.italic = italic
  if color:
    run.font.color.rgb = RGBColor.from_string(color)
  r_pr = run._element.get_or_add_rPr()
  r_fonts = r_pr.rFonts
  if r_fonts is None:
    r_fonts = OxmlElement("w:rFonts")
    r_pr.append(r_fonts)
  for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
    r_fonts.set(qn(attr), font)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
  pf = paragraph.paragraph_format
  pf.space_before = Pt(before)
  pf.space_after = Pt(after)
  pf.line_spacing = line


def set_paragraph_border(paragraph, color=BORDER, size="8", space="3", sides=("bottom",)):
  p_pr = paragraph._p.get_or_add_pPr()
  p_bdr = p_pr.find(qn("w:pBdr"))
  if p_bdr is None:
    p_bdr = OxmlElement("w:pBdr")
    p_pr.append(p_bdr)
  for side in sides:
    tag = f"w:{side}"
    element = p_bdr.find(qn(tag))
    if element is None:
      element = OxmlElement(tag)
      p_bdr.append(element)
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), size)
    element.set(qn("w:space"), space)
    element.set(qn("w:color"), color)


# ─────────────────────────── Cell helpers ──────────────────────────
def set_cell_shading(cell, fill):
  tc_pr = cell._tc.get_or_add_tcPr()
  shd = tc_pr.find(qn("w:shd"))
  if shd is None:
    shd = OxmlElement("w:shd")
    tc_pr.append(shd)
  shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="10", left_accent=None):
  tc_pr = cell._tc.get_or_add_tcPr()
  tc_borders = tc_pr.find(qn("w:tcBorders"))
  if tc_borders is None:
    tc_borders = OxmlElement("w:tcBorders")
    tc_pr.append(tc_borders)
  for edge in ("top", "left", "bottom", "right"):
    tag = f"w:{edge}"
    element = tc_borders.find(qn(tag))
    if element is None:
      element = OxmlElement(tag)
      tc_borders.append(element)
    # Viền trái dày (accent bar) nếu được yêu cầu
    if edge == "left" and left_accent:
      element.set(qn("w:val"), "single")
      element.set(qn("w:sz"), "24")
      element.set(qn("w:space"), "0")
      element.set(qn("w:color"), left_accent)
    else:
      element.set(qn("w:val"), "single")
      element.set(qn("w:sz"), size)
      element.set(qn("w:space"), "0")
      element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=170, bottom=120, end=170):
  tc_pr = cell._tc.get_or_add_tcPr()
  tc_mar = tc_pr.find(qn("w:tcMar"))
  if tc_mar is None:
    tc_mar = OxmlElement("w:tcMar")
    tc_pr.append(tc_mar)
  for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
    node = tc_mar.find(qn(f"w:{m}"))
    if node is None:
      node = OxmlElement(f"w:{m}")
      tc_mar.append(node)
    node.set(qn("w:w"), str(v))
    node.set(qn("w:type"), "dxa")


def set_table_width_pct(table, pct=100):
  tbl_pr = table._tbl.tblPr
  tbl_w = tbl_pr.find(qn("w:tblW"))
  if tbl_w is None:
    tbl_w = OxmlElement("w:tblW")
    tbl_pr.append(tbl_w)
  tbl_w.set(qn("w:type"), "pct")
  tbl_w.set(qn("w:w"), str(pct * 50))


def prevent_table_row_split(table):
  for row in table.rows:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
      tr_pr.append(OxmlElement("w:cantSplit"))


def _new_box(doc, fill, border_color, accent=None):
  """Tạo một box 1x1 (table) và trả về cell bên trong."""
  table = doc.add_table(rows=1, cols=1)
  table.alignment = WD_TABLE_ALIGNMENT.CENTER
  set_table_width_pct(table, 100)
  prevent_table_row_split(table)
  cell = table.cell(0, 0)
  set_cell_shading(cell, fill)
  set_cell_border(cell, color=border_color, size="8", left_accent=accent)
  set_cell_margins(cell)
  cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
  return cell


# ─────────────────────── Flowing (non-box) content ─────────────────
def add_body_paragraph(container, text, size=BODY_SIZE, color=TEXT_DARK,
            justify=True, after=6, first_para=False):
  """Đoạn văn thường — KHÔNG box. Đây là nơi chứa phần lớn nội dung."""
  text = strip_leading_marker(text)
  if not text.strip():
    return
  p = container.add_paragraph()
  if justify:
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
  set_paragraph_spacing(p, after=after, line=1.25)
  add_inline_runs(p, text.strip(), size=size, color=color)


def add_subheading(container, text):
  p = container.add_paragraph()
  set_paragraph_spacing(p, before=8, after=3)
  r = p.add_run(strip_markdown(text))
  set_run_font(r, font=FONT_HEAD, size=12.5, color=BRAND_NAVY, bold=True)


def add_bullet_list(container, items):
  for it in items:
    it = strip_leading_marker(it)
    if not it.strip():
      continue
    p = container.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, after=3, line=1.2)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    r = p.add_run("• ")
    set_run_font(r, size=BODY_SIZE, color=BRAND_ORANGE, bold=True)
    add_inline_runs(p, it.strip(), size=BODY_SIZE)


def strip_leading_marker(text):
  """Bỏ dấu đầu dòng markdown (-, *, +, •, số thứ tự markdown) ở đầu chuỗi."""
  if not text:
    return ""
  return re.sub(r'^\s*(?:[\*\-\+•]|\#+)\s*', '', text)


# ─────────────────────── Markdown → flowing blocks ─────────────────
KEY_PREFIX = re.compile(r'^\s*(ghi nhớ|kết luận|lưu ý|định nghĩa|quan trọng|chú ý|nhớ rằng)\s*[:：]?',
            re.IGNORECASE)


def add_markdown_body(container, raw):
  """Bộ phân tích markdown thô -> đoạn văn / danh sách / tiểu mục.

  Bảo đảm nội dung dài KHÔNG bị nhồi vào một khối duy nhất mà được tách
  thành từng đoạn tự nhiên. Dòng 'Ghi nhớ / Kết luận…' được tách ra thành
  box điểm nhấn (chỉ nội dung chính mới vào box).
  """
  if not raw:
    return
  raw = normalize_ws(raw)
  lines = raw.split('\n')
  bullet_buffer = []

  def flush_bullets():
    if bullet_buffer:
      add_bullet_list(container, list(bullet_buffer))
      bullet_buffer.clear()

  for line in lines:
    s = line.strip()
    if not s:
      flush_bullets()
      continue
    # Tiêu đề markdown -> tiểu mục
    if s.startswith('#'):
      flush_bullets()
      add_subheading(container, s)
      continue
    # Gạch đầu dòng
    if re.match(r'^\s*(?:[\*\-\+•])\s+', line) or re.match(r'^\s*\d+[\.\)]\s+', line):
      bullet_buffer.append(strip_leading_marker(s))
      continue
    flush_bullets()
    # Câu "Ghi nhớ / Kết luận" -> điểm nhấn nhỏ
    if KEY_PREFIX.match(s) and len(s) < 320:
      label = KEY_PREFIX.match(s).group(1).upper()
      body = KEY_PREFIX.sub('', s).strip()
      add_keypoint_box(container, label, body)
      continue
    add_body_paragraph(container, s)
  flush_bullets()


# ─────────────────────────── Box renderers ─────────────────────────
def add_keypoint_box(doc, label, text):
  """Box điểm nhấn NHỎ, chỉ dành cho nội dung cốt lõi (định nghĩa, ghi nhớ)."""
  text = strip_markdown_soft(text)
  cell = _new_box(doc, LIGHT_BLUE, BRAND_NAVY, accent=BRAND_NAVY)

  p = cell.paragraphs[0]
  set_paragraph_spacing(p, after=3)
  r = p.add_run(f"📌 {strip_markdown(label) or 'GHI NHỚ'}")
  set_run_font(r, font=FONT_HEAD, size=9.5, color=BRAND_NAVY, bold=True)

  p2 = cell.add_paragraph()
  p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
  set_paragraph_spacing(p2, after=0, line=1.2)
  add_inline_runs(p2, text, size=BODY_SIZE, color=TEXT_DARK)
  doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_example_box(doc, label, problem, solution):
  cell = _new_box(doc, LIGHT_ORANGE, BRAND_ORANGE, accent=BRAND_ORANGE)

  p = cell.paragraphs[0]
  set_paragraph_spacing(p, after=3)
  r = p.add_run(f"✏️ {strip_markdown(label).upper() or 'VÍ DỤ'}")
  set_run_font(r, font=FONT_HEAD, size=9.5, color=BRAND_ORANGE, bold=True)

  if problem:
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p2, after=3, line=1.2)
    add_inline_runs(p2, strip_markdown_soft(problem), size=BODY_SIZE, base_bold=True)

  if solution and solution.strip():
    p3 = cell.add_paragraph()
    set_paragraph_spacing(p3, after=0, line=1.2)
    r3 = p3.add_run("Lời giải: ")
    set_run_font(r3, font=FONT_HEAD, size=10, color=BRAND_ORANGE, italic=True, bold=True)
    add_inline_runs(p3, strip_markdown_soft(solution), size=11.5, color=TEXT_DARK)
  doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_exercise_box(doc, label, question, lines=3):
  cell = _new_box(doc, LIGHT_GRAY, BORDER, accent=BRAND_NAVY)

  p = cell.paragraphs[0]
  set_paragraph_spacing(p, after=3)
  r = p.add_run(f"📝 {strip_markdown(label).upper() or 'BÀI TẬP'}")
  set_run_font(r, font=FONT_HEAD, size=9.5, color=BRAND_NAVY, bold=True)

  q_lines = [l for l in strip_markdown_soft(question).split('\n') if l.strip()]
  for ql in (q_lines or ['']):
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p2, after=4, line=1.2)
    add_inline_runs(p2, ql.strip(), size=BODY_SIZE)

  if lines and lines > 0:
    lbl = cell.add_paragraph()
    set_paragraph_spacing(lbl, before=2, after=2)
    rl = lbl.add_run("Bài làm:")
    set_run_font(rl, font=FONT_HEAD, size=9.5, color=TEXT_MUTED, bold=True)
    for _ in range(int(lines)):
      add_dotted_line(cell)
  doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_dotted_line(cell):
  """Một dòng chấm '.....' kéo dài HẾT chiều ngang của khung (dùng dot-leader
  của Word: tab canh phải + dẫn chấm) -> luôn khít mép phải dù cỡ chữ đổi."""
  p = cell.add_paragraph()
  set_paragraph_spacing(p, after=3, line=1.0)
  tabs = p.paragraph_format.tab_stops
  tabs.add_tab_stop(Cm(CONTENT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
  r = p.add_run("\t")
  set_run_font(r, size=11, color="B9C4D4")


def strip_markdown_soft(text):
  """Bỏ heading nhưng GIỮ **đậm**/*nghiêng* để add_inline_runs biên dịch."""
  if not text:
    return ""
  text = normalize_ws(text)
  text = re.sub(r'#+\s*', '', text)
  text = text.replace('`', '')
  return text.strip()


def add_solution(doc, title, content):
  """Một mục trong phần 'Bài giải chi tiết' (đáp án): tiêu đề đậm + lời giải chảy."""
  content = strip_markdown_soft(content)
  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
  set_paragraph_spacing(p, before=4, after=2, line=1.25)
  p.paragraph_format.left_indent = Cm(0.3)
  r = p.add_run(f"{strip_markdown(title) or 'Bài'}. ")
  set_run_font(r, font=FONT_HEAD, size=BODY_SIZE, color=BRAND_ORANGE, bold=True)
  # Lời giải có thể nhiều dòng -> dòng đầu nối tiếp tiêu đề, các dòng sau thành đoạn.
  lines = [l for l in content.split('\n') if l.strip()]
  if lines:
    add_inline_runs(p, lines[0].strip(), size=BODY_SIZE)
    for extra in lines[1:]:
      pe = doc.add_paragraph()
      pe.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
      set_paragraph_spacing(pe, after=2, line=1.25)
      pe.paragraph_format.left_indent = Cm(0.3)
      add_inline_runs(pe, extra.strip(), size=BODY_SIZE)


def add_image(doc, block):
  path = block.get('path')
  caption = block.get('caption') or block.get('desc') or ''
  if path and os.path.exists(path):
    try:
      p = doc.add_paragraph()
      p.alignment = WD_ALIGN_PARAGRAPH.CENTER
      set_paragraph_spacing(p, before=4, after=2)
      run = p.add_run()
      run.add_picture(path, width=Inches(4.3))
      if caption:
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(pc, after=8)
        rc = pc.add_run(f"Hình: {clean_caption(caption)}")
        set_run_font(rc, font=FONT_HEAD, size=9.5, color=TEXT_MUTED, italic=True)
      return
    except Exception:
      pass
  # Không có ảnh -> placeholder gọn gàng (báo hiệu cần image-fetcher/Artist)
  cell = _new_box(doc, WHITE, BORDER)
  p = cell.paragraphs[0]
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  r = p.add_run(f"🖼 [Chèn hình: {clean_caption(caption) or 'minh họa'}]")
  set_run_font(r, size=10.5, color=TEXT_MUTED, italic=True)
  doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ─────────────────────────── Titles ────────────────────────────────
def add_document_title(doc, title, subtitle=""):
  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  r = p.add_run("KI-TEE EDUCATION")
  set_run_font(r, font=FONT_HEAD, size=10, color=BRAND_ORANGE, bold=True)
  set_paragraph_spacing(p, before=4, after=4)

  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  r = p.add_run(strip_markdown(title).upper())
  set_run_font(r, font=FONT_HEAD, size=20, color=BRAND_NAVY, bold=True)
  set_paragraph_spacing(p, before=0, after=2)

  if subtitle:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(strip_markdown(subtitle))
    set_run_font(r, font=FONT_HEAD, size=10.5, color=TEXT_MUTED)
    set_paragraph_spacing(p, before=0, after=8)

  line = doc.add_paragraph()
  set_paragraph_border(line, color=BRAND_ORANGE, size="12", sides=("bottom",))
  set_paragraph_spacing(line, before=0, after=10)


def add_section_title(doc, title):
  p = doc.add_paragraph()
  r = p.add_run(strip_markdown(title))
  set_run_font(r, font=FONT_HEAD, size=14, color=BRAND_NAVY, bold=True)
  set_paragraph_spacing(p, before=12, after=4)
  set_paragraph_border(p, color=BORDER, size="6", sides=("bottom",))


# ─────────────────────────── Dispatch ──────────────────────────────
def render_block(doc, block):
  t = (block.get('type') or 'paragraph').lower()

  if t in ('subheading', 'heading'):
    add_subheading(doc, block.get('text') or block.get('title', ''))

  elif t == 'paragraph':
    add_markdown_body(doc, block.get('text') or block.get('content', ''))

  elif t == 'list':
    items = block.get('items')
    if items:
      add_bullet_list(doc, items)
    else:
      add_markdown_body(doc, block.get('content', ''))

  elif t in ('keypoint', 'note'):
    add_keypoint_box(doc, block.get('title', 'GHI NHỚ'),
             block.get('text') or block.get('content', ''))

  elif t == 'theory':
    # Nội dung lý thuyết (thường dài): tách thành đoạn văn thường,
    # chỉ câu "Ghi nhớ/Kết luận" mới thành box (xử lý trong add_markdown_body).
    title = block.get('title')
    if title and title.lower() not in ('lý thuyết', 'khái niệm'):
      add_subheading(doc, title)
    add_markdown_body(doc, block.get('content') or block.get('text', ''))

  elif t == 'example':
    add_example_box(doc, block.get('title', 'Ví dụ'),
            block.get('problem') or block.get('content', ''),
            block.get('solution', ''))

  elif t == 'exercise':
    add_exercise_box(doc, block.get('title', 'Bài tập'),
             block.get('question') or block.get('content', ''),
             block.get('lines', block.get('answerLines', 3)))

  elif t == 'solution':
    add_solution(doc, block.get('title', 'Bài'),
           block.get('content') or block.get('solution', ''))

  elif t in ('image', 'figure'):
    add_image(doc, block)

  else:
    add_markdown_body(doc, block.get('content') or block.get('text', ''))


def create_styled_word(doc_model, output_path):
  global CONTENT_WIDTH_CM
  doc = Document()
  section = doc.sections[0]
  section.page_width = Cm(21.0)   # A4
  section.page_height = Cm(29.7)
  section.top_margin = Cm(1.5)
  section.bottom_margin = Cm(1.5)
  section.left_margin = Cm(1.7)
  section.right_margin = Cm(1.7)
  # Chiều rộng trong khung = trang - lề - lề trong ô (~0.6cm) - đệm nhỏ
  CONTENT_WIDTH_CM = 21.0 - 1.7 - 1.7 - 0.6 - 0.2

  add_document_title(
    doc,
    doc_model.get('title', 'Tài liệu học tập'),
    f"Môn: {doc_model.get('subject', '')} | Lớp: {doc_model.get('grade', '')}"
  )

  for section_data in doc_model.get('sections', []):
    add_section_title(doc, section_data.get('heading', ''))
    for block in section_data.get('blocks', []):
      render_block(doc, block)

  doc.save(output_path)


if __name__ == "__main__":
  with open(sys.argv[1], 'r', encoding='utf-8') as f:
    model = json.load(f)
  create_styled_word(model, sys.argv[2])
  print(f"✅ Word saved: {sys.argv[2]}")
